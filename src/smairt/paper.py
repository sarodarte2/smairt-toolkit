"""Paper-element provenance validation."""

from __future__ import annotations

import hashlib
import json
import tempfile
import zipfile
from pathlib import Path
from typing import Any, cast
from uuid import uuid4

import yaml

from smairt.integrity import verify_run
from smairt.locking import mutating
from smairt.models import (
    ClaimRecord,
    ClaimStatus,
    EvidenceCard,
    EvidenceStatus,
    RunRecord,
    RunStatus,
    utc_now,
)
from smairt.provenance import record_event, require_contributor, stage_event
from smairt.references import load_index
from smairt.transactions import FileTransaction
from smairt.utils import (
    atomic_write,
    ensure_no_symlink,
    sha256_file,
    sha256_text,
    slugify,
    validate_identifier,
)


def accepted_runs(root: Path) -> dict[str, dict[str, Any]]:
    """Return only completed, relationship-consistent, integrity-verified selections."""
    accepted: dict[str, dict[str, Any]] = {}
    for selection in (root / "analysis").glob("*/selection.yaml"):
        payload = yaml.safe_load(selection.read_text()) or {}
        if not isinstance(payload, dict) or payload.get("status") != "ACCEPTED":
            continue
        experiment_id = str(payload.get("experiment_id", selection.parent.name))
        iteration_id = str(payload.get("iteration_id", ""))
        run_id = str(payload.get("run_id", ""))
        for label, identifier in (
            ("experiment ID", experiment_id),
            ("iteration ID", iteration_id),
            ("run ID", run_id),
        ):
            validate_identifier(identifier, label=label)
        if experiment_id != selection.parent.name:
            raise ValueError("accepted selection does not match its analysis directory")
        run_path = root / "results" / experiment_id / iteration_id / run_id / "run.json"
        record = RunRecord.model_validate_json(run_path.read_text(encoding="utf-8"))
        if (
            record.run_id != run_id
            or record.experiment_id != experiment_id
            or record.iteration_id != iteration_id
            or record.status is not RunStatus.COMPLETED
            or record.exit_code != 0
        ):
            raise ValueError(f"accepted run is not a completed matching execution: {run_id}")
        if not verify_run(root, run_id)["ok"]:
            raise ValueError(f"accepted run failed integrity verification: {run_id}")
        if run_id in accepted:
            raise ValueError(f"duplicate accepted run selection: {run_id}")
        accepted[run_id] = {**payload, "run_path": str(run_path.relative_to(root))}
    return accepted


def validate_paper(root: Path) -> list[str]:
    """Ensure every paper element points to unique, currently accepted evidence."""
    manifest = root / "paper/manifest.yaml"
    payload = yaml.safe_load(manifest.read_text()) or {}
    elements = payload.get("elements", [])
    if not isinstance(elements, list):
        return ["paper manifest elements must be a list"]
    errors: list[str] = []
    try:
        accepted = accepted_runs(root)
    except (OSError, TypeError, ValueError) as exc:
        accepted = {}
        errors.append(f"invalid accepted-run state: {exc}")
    identifiers: set[str] = set()
    for index, element in enumerate(elements, start=1):
        if not isinstance(element, dict):
            errors.append(f"paper element {index} must be a mapping")
            continue
        identifier = str(element.get("id", ""))
        if not identifier:
            errors.append(f"paper element {index} is missing id")
        elif identifier in identifiers:
            errors.append(f"duplicate paper element id: {identifier}")
        identifiers.add(identifier)
        run_id = str(element.get("run_id", ""))
        try:
            validate_identifier(run_id, label="run ID")
        except ValueError:
            errors.append(f"{identifier or index} has an invalid run ID")
            continue
        if run_id not in accepted:
            display_run = run_id or "[missing]"
            errors.append(f"{identifier or index} references non-accepted run {display_run}")
    for claim_path in sorted((root / "paper/claims").glob("*.json")):
        try:
            claim = ClaimRecord.model_validate_json(claim_path.read_text(encoding="utf-8"))
            if claim.id != claim_path.stem:
                raise ValueError("claim ID does not match its filename")
        except (OSError, ValueError) as exc:
            errors.append(f"invalid claim record {claim_path.name}: {exc}")
            continue
        if claim.status is not ClaimStatus.APPROVED:
            continue
        for evidence_id in claim.evidence_ids:
            evidence_path = root / "paper/evidence" / f"{evidence_id}.json"
            if not evidence_path.exists():
                errors.append(f"approved claim {claim.id} has missing evidence {evidence_id}")
                continue
            try:
                evidence = EvidenceCard.model_validate_json(
                    evidence_path.read_text(encoding="utf-8")
                )
                if evidence.id != evidence_id:
                    raise ValueError("evidence ID does not match its filename")
                selection = accepted[evidence.run_id]
                run_path = root / str(selection["run_path"])
                if evidence.run_record_sha256 != sha256_file(run_path):
                    raise ValueError("evidence run-record hash is stale")
            except (KeyError, OSError, ValueError) as exc:
                errors.append(f"approved claim {claim.id} has stale evidence {evidence_id}: {exc}")
                continue
            if evidence.status is not EvidenceStatus.CURRENT:
                errors.append(f"approved claim {claim.id} has stale evidence {evidence_id}")
        references = {item.id: item for item in load_index(root)}
        for reference_id in claim.reference_ids:
            if (
                reference_id not in references
                or references[reference_id].verification_status != "verified"
            ):
                errors.append(f"approved claim {claim.id} has unverified citation {reference_id}")
    return errors


@mutating("paper evidence create")
def create_evidence_card(
    root: Path,
    run_id: str,
    *,
    purpose: str,
    observed_result: str,
    limitations: str,
    decision: str,
    relevance: str = "",
) -> Path:
    """Freeze one accepted run into a contributor-attributed evidence card."""
    validate_identifier(run_id, label="run ID")
    contributor = require_contributor(root)
    accepted = accepted_runs(root)
    if run_id not in accepted:
        raise ValueError("evidence cards require a currently accepted run")
    verification = verify_run(root, run_id)
    if not verification["ok"]:
        raise ValueError("run integrity verification failed")
    run_path = root / str(accepted[run_id]["run_path"])
    payload = {
        "schema_version": 1,
        "id": f"evidence-{run_id.lower()}",
        "run_id": run_id,
        "purpose": purpose,
        "observed_result": observed_result,
        "limitations": limitations,
        "decision": decision,
        "possible_paper_relevance": relevance,
        "contributor": contributor.id,
        "created_at": utc_now(),
        "run_record_sha256": sha256_file(run_path),
        "status": "current",
    }
    card = EvidenceCard.model_validate(payload)
    path = root / "paper/evidence" / f"{card.id}.json"
    if path.exists():
        raise ValueError(f"evidence card already exists for {run_id}")
    transaction = FileTransaction(root, "paper evidence create")
    rendered = json.dumps(card.model_dump(mode="json"), indent=2, sort_keys=True) + "\n"
    transaction.stage_text(path, rendered)
    stage_event(
        root,
        transaction,
        "paper.evidence.created",
        artifact_ids=[card.id],
        hashes={str(path.relative_to(root)): sha256_text(rendered)},
    )
    transaction.commit()
    return path


@mutating("paper claim create")
def create_claim(
    root: Path, statement: str, evidence_ids: list[str], reference_ids: list[str] | None = None
) -> Path:
    """Propose a uniquely identified claim linked to existing evidence cards."""
    contributor = require_contributor(root)
    for identifier in evidence_ids:
        validate_identifier(identifier, label="evidence ID")
        evidence_path = root / "paper/evidence" / f"{identifier}.json"
        if not evidence_path.exists():
            raise ValueError(f"unknown evidence card: {identifier}")
        EvidenceCard.model_validate_json(evidence_path.read_text(encoding="utf-8"))
    for identifier in reference_ids or []:
        validate_identifier(identifier, label="reference ID")
    identifier = f"claim-{slugify(statement)[:40]}-{uuid4().hex[:8]}"
    payload = {
        "schema_version": 1,
        "id": identifier,
        "statement": statement,
        "evidence_ids": evidence_ids,
        "reference_ids": reference_ids or [],
        "status": "proposed",
        "proposed_by": contributor.id,
        "created_at": utc_now(),
    }
    claim = ClaimRecord.model_validate(payload)
    path = root / "paper/claims" / f"{identifier}.json"
    transaction = FileTransaction(root, "paper claim create")
    transaction.stage_text(
        path, json.dumps(claim.model_dump(mode="json"), indent=2, sort_keys=True) + "\n"
    )
    stage_event(root, transaction, "paper.claim.proposed", artifact_ids=[identifier])
    transaction.commit()
    return path


@mutating("paper claim review")
def review_claim(root: Path, identifier: str, status: str) -> dict[str, Any]:
    """Record a human claim decision after validating approval dependencies."""
    validate_identifier(identifier, label="claim ID")
    if status not in {"approved", "rejected", "retracted", "superseded"}:
        raise ValueError("invalid claim state")
    contributor = require_contributor(root)
    path = root / "paper/claims" / f"{identifier}.json"
    claim = ClaimRecord.model_validate_json(path.read_text(encoding="utf-8"))
    if claim.id != identifier:
        raise ValueError("claim ID does not match its filename")
    payload = claim.model_dump(mode="json")
    if status == "approved":
        errors = _validate_claim(root, payload)
        if errors:
            raise ValueError("claim cannot be approved: " + "; ".join(errors))
    payload.update(status=status, reviewed_by=contributor.id, reviewed_at=utc_now())
    reviewed = ClaimRecord.model_validate(payload)
    transaction = FileTransaction(root, "paper claim review")
    transaction.stage_text(
        path, json.dumps(reviewed.model_dump(mode="json"), indent=2, sort_keys=True) + "\n"
    )
    stage_event(root, transaction, f"paper.claim.{status}", artifact_ids=[identifier])
    transaction.commit()
    return reviewed.model_dump(mode="json")


@mutating("paper begin")
def begin_paper(root: Path, title: str) -> Path:
    """Create the authoritative manuscript only after approved claims exist."""
    claims = [json.loads(path.read_text()) for path in (root / "paper/claims").glob("*.json")]
    approved = [claim for claim in claims if claim.get("status") == "approved"]
    if not approved:
        raise ValueError("paper drafting requires at least one approved claim")
    path = root / "paper/manuscript.md"
    if path.exists():
        raise FileExistsError(path)
    transaction = FileTransaction(root, "paper begin")
    transaction.stage_text(
        path,
        f"# {title}\n\n## Abstract\n\n## Introduction\n\n## Methods\n\n"
        "## Results\n\n## Discussion\n\n## References\n",
    )
    transaction.stage_text(
        root / "paper/sections.yaml",
        yaml.safe_dump(
            {
                "sections": {
                    name: {"status": "draft", "claim_ids": []}
                    for name in (
                        "Abstract",
                        "Introduction",
                        "Methods",
                        "Results",
                        "Discussion",
                        "References",
                    )
                }
            },
            sort_keys=False,
        ),
    )
    transaction.commit()
    record_event(root, "paper.begun", artifact_ids=["paper/manuscript.md"])
    return path


@mutating("paper outline create")
def create_outline(root: Path) -> Path:
    """Create a claim-linked outline from approved claims without drafting prose."""
    claims = [json.loads(path.read_text()) for path in (root / "paper/claims").glob("*.json")]
    approved = [claim for claim in claims if claim.get("status") == "approved"]
    if not approved:
        raise ValueError("an outline requires at least one approved claim")
    lines = ["# Manuscript Outline", ""]
    for section in ("Introduction", "Methods", "Results", "Discussion"):
        lines.extend([f"## {section}", ""])
        lines.extend(f"- [{claim['id']}] {claim['statement']}" for claim in approved)
        lines.append("")
    path = root / "paper/outline.md"
    atomic_write(path, "\n".join(lines))
    record_event(root, "paper.outline.created", artifact_ids=["paper/outline.md"])
    return path


@mutating("paper section draft")
def draft_section(root: Path, section: str, text: str, claim_ids: list[str]) -> Path:
    """Replace one canonical Markdown section with prose linked to approved claims."""
    manuscript = root / "paper/manuscript.md"
    if not manuscript.exists():
        raise ValueError("paper has not begun")
    claims = {p.stem: json.loads(p.read_text()) for p in (root / "paper/claims").glob("*.json")}
    invalid = [
        identifier
        for identifier in claim_ids
        if identifier not in claims or claims[identifier].get("status") != "approved"
    ]
    if invalid:
        raise ValueError("section references non-approved claims: " + ", ".join(invalid))
    content = manuscript.read_text()
    heading = f"## {section}"
    if heading not in content:
        raise ValueError(f"unknown manuscript section: {section}")
    before, remainder = content.split(heading, 1)
    next_heading = remainder.find("\n## ")
    after = remainder[next_heading:] if next_heading >= 0 else "\n"
    linked = " ".join(f"[{identifier}]" for identifier in claim_ids)
    manuscript_content = (
        f"{before}{heading}\n\n{text.strip()}\n\nClaims: {linked}\n{after.lstrip()}"
    )
    sections_path = root / "paper/sections.yaml"
    payload = yaml.safe_load(sections_path.read_text()) or {"sections": {}}
    payload.setdefault("sections", {}).setdefault(section, {})
    payload["sections"][section].update(status="draft", claim_ids=claim_ids, drafted_at=utc_now())
    transaction = FileTransaction(root, "paper section draft")
    transaction.stage_text(manuscript, manuscript_content)
    transaction.stage_text(sections_path, yaml.safe_dump(payload, sort_keys=False))
    transaction.commit()
    record_event(root, "paper.section.drafted", artifact_ids=[section])
    return manuscript


@mutating("paper section review")
def review_section(root: Path, section: str, claim_ids: list[str]) -> dict[str, Any]:
    """Mark a manuscript section reviewed against approved supporting claims."""
    require_contributor(root)
    claims = {p.stem: json.loads(p.read_text()) for p in (root / "paper/claims").glob("*.json")}
    invalid = [
        identifier
        for identifier in claim_ids
        if identifier not in claims or claims[identifier].get("status") != "approved"
    ]
    if invalid:
        raise ValueError("section references non-approved claims: " + ", ".join(invalid))
    path = root / "paper/sections.yaml"
    payload = yaml.safe_load(path.read_text()) or {"sections": {}}
    payload.setdefault("sections", {}).setdefault(section, {})
    payload["sections"][section].update(
        status="reviewed", claim_ids=claim_ids, reviewed_at=utc_now()
    )
    atomic_write(path, yaml.safe_dump(payload, sort_keys=False))
    record_event(root, "paper.section.reviewed", artifact_ids=[section])
    return cast(dict[str, Any], payload["sections"][section])


@mutating("paper build")
def build_paper(
    root: Path,
    format_name: str,
    *,
    template: Path | None = None,
    line_numbering: bool = False,
) -> Path:
    """Create a versioned snapshot from canonical Markdown after provenance validation."""
    errors = validate_paper(root)
    if errors:
        raise ValueError("paper validation failed: " + "; ".join(errors))
    manuscript = root / "paper/manuscript.md"
    if not manuscript.exists():
        raise ValueError("paper has not begun")
    sections_path = root / "paper/sections.yaml"
    sections = (yaml.safe_load(sections_path.read_text()) or {}).get("sections", {})
    unreviewed = [name for name, value in sections.items() if value.get("status") != "reviewed"]
    if not sections or unreviewed:
        raise ValueError("paper sections require explicit review: " + ", ".join(unreviewed))
    stamp = utc_now().replace(":", "").replace("+", "_")
    manuscript_content = manuscript.read_text(encoding="utf-8")
    if format_name == "md":
        target = root / "paper/builds" / f"manuscript-{stamp}.md"
        output = manuscript_content.encode("utf-8")
    elif format_name == "docx":
        target = root / "paper/builds" / f"manuscript-{stamp}.docx"
        template = _validated_docx_template(template)
        with tempfile.TemporaryDirectory(prefix="smairt-docx-", dir=root / ".smairt") as temporary:
            staged = Path(temporary) / "manuscript.docx"
            _write_docx(
                staged,
                manuscript_content,
                template=template,
                line_numbering=line_numbering,
                paper_root=root / "paper",
            )
            output = staged.read_bytes()
    else:
        raise ValueError("format must be md or docx")
    transaction = FileTransaction(root, "paper build")
    transaction.stage_bytes(target, output)
    stage_event(
        root,
        transaction,
        "paper.built",
        artifact_ids=[str(target.relative_to(root))],
        hashes={str(target.relative_to(root)): hashlib.sha256(output).hexdigest()},
    )
    transaction.commit()
    return target


def _validated_docx_template(template: Path | None) -> Path | None:
    """Reject unsafe, oversized, or structurally invalid DOCX templates."""
    if template is None:
        return None
    requested = template.expanduser().absolute()
    if requested.is_symlink() or any(parent.is_symlink() for parent in requested.parents):
        raise ValueError("DOCX template must not traverse symlinks")
    if not requested.is_file() or requested.suffix.lower() != ".docx":
        raise ValueError("DOCX template must be an existing .docx file")
    if requested.stat().st_size > 50 * 1024 * 1024 or not zipfile.is_zipfile(requested):
        raise ValueError("DOCX template is oversized or structurally invalid")
    return requested.resolve()


def _write_docx(
    path: Path,
    markdown: str,
    *,
    template: Path | None = None,
    line_numbering: bool = False,
    paper_root: Path,
) -> None:
    """Write a journal-neutral DOCX using Word-native styles and page structure."""
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.shared import Inches, Pt

    document = Document(str(template)) if template else Document()
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    lines = markdown.splitlines()
    title = next((line[2:].strip() for line in lines if line.startswith("# ")), "Manuscript")
    title_paragraph = document.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(18)
    document.add_paragraph("Authors and affiliations: complete manually")
    document.add_paragraph("Corresponding author: complete manually")
    document.add_section(WD_SECTION.NEW_PAGE)
    if line_numbering:
        for section in document.sections:
            line_numbers = OxmlElement("w:lnNumType")
            line_numbers.set(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}countBy", "1"
            )
            section._sectPr.append(line_numbers)
    for line in lines:
        if not line.strip() or line.startswith("# "):
            continue
        if line.startswith("## "):
            document.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            document.add_heading(line[4:].strip(), level=2)
        elif line.startswith("![") and "](" in line:
            alt, source = line[2:].split("](", 1)
            resolved_paper_root = paper_root.resolve()
            image = ensure_no_symlink(resolved_paper_root, resolved_paper_root / source.rstrip(")"))
            if image.exists():
                document.add_picture(str(image), width=Inches(6))
                document.add_paragraph(alt, style="Caption")
        else:
            document.add_paragraph(line)
    path.parent.mkdir(parents=True, exist_ok=True)
    # python-docx writes incrementally, so render outside the final filename and
    # publish only after the document package closes successfully.
    with tempfile.NamedTemporaryFile(
        prefix=f".{path.name}.", suffix=".docx", dir=path.parent, delete=False
    ) as stream:
        temporary = Path(stream.name)
    try:
        document.save(str(temporary))
        temporary.replace(path)
    finally:
        temporary.unlink(missing_ok=True)


def _validate_claim(root: Path, claim: dict[str, Any]) -> list[str]:
    """Return evidence and citation errors that prevent claim approval."""
    errors: list[str] = []
    try:
        record = ClaimRecord.model_validate(claim)
    except ValueError as exc:
        return [f"invalid claim record: {exc}"]
    try:
        accepted = accepted_runs(root)
    except (OSError, TypeError, ValueError) as exc:
        accepted = {}
        errors.append(f"invalid accepted-run state: {exc}")
    for evidence_id in record.evidence_ids:
        path = root / "paper/evidence" / f"{evidence_id}.json"
        if not path.exists():
            errors.append(f"missing evidence {evidence_id}")
            continue
        try:
            evidence = EvidenceCard.model_validate_json(path.read_text(encoding="utf-8"))
            selection = accepted[evidence.run_id]
            run_path = root / str(selection["run_path"])
            if evidence.id != evidence_id or evidence.run_record_sha256 != sha256_file(run_path):
                raise ValueError("evidence identity or run hash is stale")
        except (KeyError, OSError, ValueError) as exc:
            errors.append(f"stale evidence {evidence_id}: {exc}")
            continue
        if evidence.status is not EvidenceStatus.CURRENT:
            errors.append(f"stale evidence {evidence_id}")
    try:
        references = {item.id: item for item in load_index(root)}
    except (OSError, TypeError, ValueError) as exc:
        references = {}
        errors.append(f"invalid reference index: {exc}")
    for reference_id in record.reference_ids:
        if (
            reference_id not in references
            or references[reference_id].verification_status != "verified"
        ):
            errors.append(f"unverified citation {reference_id}")
    return errors
